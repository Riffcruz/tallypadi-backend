from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "TallyPadi_Marketer_Feature_Guide.docx"

NAIRA = "\u20a6"

COLORS = {
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "ink": "0B2545",
    "muted": "5B677A",
    "light_fill": "E8EEF5",
    "pale": "F4F6F9",
    "green": "0F766E",
    "gold": "7A5A00",
    "red": "9B1C1C",
    "border": "C7D0DD",
    "white": "FFFFFF",
}


def rgb(hex_color: str) -> RGBColor:
    hex_color = hex_color.replace("#", "")
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_in):
    width = int(width_in * 1440)
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color=COLORS["border"], size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, total_width_in=6.5):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(total_width_in * 1440)))
    tbl_w.set(qn("w:type"), "dxa")


def clear_cell(cell):
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)


def write_cell(cell, text, bold=False, color="0B2545", size=9.3):
    clear_cell(cell)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    for idx, part in enumerate(str(text).split("\n")):
        if idx:
            p.add_run().add_break()
        run = p.add_run(part)
        set_run_font(run, size=size, color=color, bold=bold)


def add_table(doc, headers, rows, widths, header_fill=COLORS["light_fill"]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    set_table_borders(table)
    header_cells = table.rows[0].cells
    for idx, label in enumerate(headers):
        set_cell_width(header_cells[idx], widths[idx])
        set_cell_margins(header_cells[idx])
        set_cell_fill(header_cells[idx], header_fill)
        header_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        write_cell(header_cells[idx], label, bold=True, color=COLORS["ink"], size=9.2)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_width(cells[idx], widths[idx])
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            write_cell(cells[idx], value, color=COLORS["ink"], size=9.0)
    doc.add_paragraph()
    return table


def add_para(doc, text="", size=10.5, color=COLORS["ink"], bold=False, italic=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        set_run_font(run, size=10.2, color=COLORS["ink"])


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        set_run_font(run, size=10.2, color=COLORS["ink"])


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_callout(doc, title, body, fill=COLORS["pale"], accent=COLORS["green"]):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    set_table_borders(table, color="D7DEE8", size="4")
    cell = table.rows[0].cells[0]
    set_cell_fill(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    clear_cell(cell)
    p1 = cell.add_paragraph()
    p1.paragraph_format.space_after = Pt(3)
    r1 = p1.add_run(title)
    set_run_font(r1, size=10.4, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(body)
    set_run_font(r2, size=9.6, color=COLORS["ink"])
    doc.add_paragraph()


def setup_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(COLORS["ink"])
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.bold = True
        style.font.color.rgb = rgb(COLORS["blue"] if style_name != "Heading 3" else COLORS["dark_blue"])
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.25

    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].paragraph_format.space_before = Pt(18)
    styles["Heading 1"].paragraph_format.space_after = Pt(10)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].paragraph_format.space_before = Pt(14)
    styles["Heading 2"].paragraph_format.space_after = Pt(7)
    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 3"].paragraph_format.space_before = Pt(10)
    styles["Heading 3"].paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.text = ""
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("TallyPadi Marketer Feature Guide")
    set_run_font(run, size=8.5, color=COLORS["muted"], bold=True)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Prepared for marketer enablement")
    set_run_font(run, size=8, color=COLORS["muted"])


def add_cover(doc):
    add_para(doc, "TALLYPADI", size=12, color=COLORS["green"], bold=True, after=12)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.line_spacing = 1.05
    run = title.add_run("Marketer Feature Guide")
    set_run_font(run, size=28, color=COLORS["ink"], bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    r = subtitle.add_run("Everything marketers need to understand, explain, demo, and sell the TallyPadi business.")
    set_run_font(r, size=13, color=COLORS["muted"])

    add_table(
        doc,
        ["Guide Field", "Details"],
        [
            ["Audience", "Marketers, sales agents, partners, customer success, support leads."],
            ["Purpose", "Explain the business clearly, map every major feature to customer benefits, and reduce overpromising."],
            ["Product reviewed", "TallyPadi backend and web codebase in this workspace."],
            ["Last reviewed", date(2026, 5, 22).strftime("%B %d, %Y")],
        ],
        [1.5, 5.0],
    )

    add_callout(
        doc,
        "Core positioning",
        "TallyPadi is an all-in-one WhatsApp + web dashboard business manager for SMEs. "
        "It helps merchants record sales, control stock, send receipts and invoices, track debts, publish an online storefront, list products in a marketplace, and request managed ads from one practical workflow.",
    )

    add_heading(doc, "How To Use This Guide", level=1)
    add_bullets(
        doc,
        [
            "Use the executive summary for quick pitches.",
            "Use the feature sections to answer customer questions accurately.",
            "Use the demo script to show the product in a convincing order.",
            "Use the limits section to avoid promising features the code does not currently implement.",
        ],
    )
    doc.add_page_break()


def add_executive_summary(doc):
    add_heading(doc, "Executive Summary", level=1)
    add_para(
        doc,
        "TallyPadi is built for business owners who want the speed of WhatsApp and the control of a proper dashboard. "
        "The product covers day-to-day shop operations, customer records, financial visibility, online selling, staff oversight, ads, and platform support.",
    )
    add_table(
        doc,
        ["Business Area", "What TallyPadi Does", "Simple Benefit For Customers"],
        [
            ["WhatsApp operations", "Parses plain-language messages for sales, restocks, reports, debt payments, invoices, orders, settings, and support.", "Owners keep records without learning complicated software."],
            ["POS and sales ledger", "Records sales from dashboard or chat, deducts stock, supports discounts, payment methods, receipts, sales history, and register close reports.", "Cashiers move faster and owners see what happened."],
            ["Inventory", "Tracks product names, SKU, barcode, category, images, cost price, selling price, stock, publication status, low-stock thresholds, and supplier details.", "Owners know what is available, what is low, and what is selling."],
            ["Customers and debtors", "Stores customers, loyalty points, credit sales, debtor balances, payments, due dates, and reminders.", "Merchants reduce lost debts and reward repeat buyers."],
            ["Online store and marketplace", "Publishes a public shop link, product catalog, marketplace listings, seller verification badge, search, filters, product pages, and visit tracking.", "Merchants can show stock online and look more trustworthy."],
            ["Ads manager", "Lets Tycoon merchants fund an ads wallet, submit product boosts, select Meta/TikTok/Google/Marketplace channels, track review status, top up, pause, stop, resume, and see metrics.", "Merchants can request managed visibility without becoming ad experts."],
            ["Team and branch controls", "Supports staff login, permissions, staff sale alerts, HQ branches, warehouses, branch comparison, and stock transfers.", "The system grows from one-person shops to multi-location operations."],
            ["Platform operations", "Provides admin analytics, support tickets, broadcasts, blog CMS, seller verification review, ads review, investor dashboards, queues, sockets, and schedulers.", "TallyPadi can operate the service and support customers at scale."],
        ],
        [1.35, 3.05, 2.1],
    )

    add_heading(doc, "Best One-Line Pitch", level=2)
    add_callout(
        doc,
        "Use this sentence",
        "TallyPadi turns WhatsApp into a simple business manager, then gives the owner a dashboard, receipts, reports, storefront, marketplace, and ads tools as the business grows.",
        fill="F0F7F5",
        accent=COLORS["green"],
    )

    add_heading(doc, "Primary Buyers", level=2)
    add_bullets(
        doc,
        [
            "Retail shop owners who currently use notebooks, memory, or scattered WhatsApp chats.",
            "Wholesalers and distributors that need stock and sales visibility.",
            "Fashion, beauty, electronics, grocery, phone accessories, furniture, and appliance sellers.",
            "Service or made-to-order businesses that need orders, invoices, deposits, and delivery dates.",
            "Growing merchants that need staff accounts, branch oversight, online catalogs, and paid promotion.",
            "Marketplace sellers who want a verified identity badge and product discovery.",
        ],
    )
    doc.add_page_break()


def add_customer_journey(doc):
    add_heading(doc, "Customer Journey", level=1)
    add_numbered(
        doc,
        [
            "Merchant signs up with phone, email, shop name, password, country, and email OTP verification.",
            "Merchant starts with a 7-day trial, then can renew through Paystack for Oga Boss or Tycoon.",
            "Merchant adds products manually, through web upload, or by chatting with the WhatsApp assistant.",
            "Merchant records sales from WhatsApp or the dashboard. Stock reduces automatically and receipts can be sent.",
            "Merchant tracks debtors, payments, expenses, customers, loyalty points, and orders.",
            "Merchant publishes a Tycoon online store with a shop link, product catalog, location, theme, and images.",
            "Merchant can submit seller verification to earn a visible marketplace verification badge.",
            "Merchant can fund an ads wallet and submit product boosts for TallyPadi review and fulfillment.",
            "As the business grows, the merchant adds staff, sets permissions, promotes HQ managers, creates branches or warehouses, and compares branch performance.",
        ],
    )
    add_heading(doc, "What A Marketer Should Emphasize", level=2)
    add_bullets(
        doc,
        [
            "The customer does not have to abandon WhatsApp. TallyPadi meets them where they already work.",
            "The dashboard is there when the merchant needs speed, reports, staff control, or a cleaner POS workflow.",
            "The store link and marketplace make inventory visible to buyers instead of locked inside a shop notebook.",
            "The ads product is managed and review-based, so merchants can promote products without managing ad accounts themselves.",
            "The strongest message is control: stock control, sales control, debtor control, staff control, and growth control.",
        ],
    )


def add_feature_sections(doc):
    sections = [
        (
            "Account, Login, Plans, And Billing",
            [
                "Owner registration with phone number, email, shop name, password, country, and email OTP verification.",
                "Owner login with email or phone number plus password.",
                "Staff login through WhatsApp OTP, with checks that the owner is active and on the correct plan.",
                "Forgot-password OTP and change-phone OTP flows using WhatsApp.",
                "Subscription states include trial, active, past due, cancelled, and suspended.",
                "Paystack subscription payments support Oga Boss and Tycoon with 1-month, 6-month, and 12-month durations.",
                f"Current pricing in the payment flow: Oga Boss {NAIRA}3,000 monthly, {NAIRA}15,000 for 6 months, {NAIRA}28,800 yearly; Tycoon {NAIRA}5,000 monthly, {NAIRA}27,000 for 6 months, {NAIRA}42,000 yearly.",
                "A scheduled job expires trials and active subscriptions when due dates pass.",
            ],
            "Pitch: Merchants can start quickly, verify ownership, and pay locally through Paystack. For marketers, always confirm final plan packaging before publishing public collateral because older landing copy and backend gates differ in a few places.",
        ),
        (
            "WhatsApp Business Assistant",
            [
                "Parses plain English, Pidgin, and other selected languages through Gemini prompts.",
                "Handles sales, restock, set stock, delete stock, price checks, price updates, debt payments, expenses, reports, orders, shop links, invoices, settings, support, and HQ commands.",
                "Uses conversation history to complete partial instructions and context such as undoing the last sale.",
                "Supports interactive WhatsApp buttons, lists, flows, CTA links, receipts, invoices, and document sending.",
                "Sends typing indicators, read receipts, queued replies, support handoff, and registration completion menus.",
                "Includes safety parsing for messy money values, quantities, item names, date ranges, due dates, credit sales, and report requests.",
            ],
            "Pitch: The merchant can type the way they speak. TallyPadi turns the message into structured business records.",
        ),
        (
            "Sales, POS, Receipts, And Register Close",
            [
                "Records sales from the web dashboard or WhatsApp.",
                "Supports multiple items, quantity, selling price, cost price snapshots, payment method, discount amount, and customer selection.",
                "Checks stock before web sales and prevents insufficient-stock sales.",
                "Deducts stock automatically and records the actual seller, including staff.",
                "Generates sale receipt PDFs and sends them through WhatsApp or downloads from the API.",
                "Supports undo/delete sale workflows and refund/return logic in the service layer.",
                "Includes close-register/Z-report workflow where staff submit physical cash and owners receive expected cash, actual cash, and discrepancy.",
                "Sends staff sale alerts to owners when enabled and the owner was recently active.",
            ],
            "Pitch: TallyPadi replaces scattered records with a clean sales ledger and professional receipts.",
        ),
        (
            "Inventory And Stock Management",
            [
                "Products include name, SKU, stock quantity, selling price, cost price, image, category, barcode, description, colors, sizes, and public-publish status.",
                "Auto-generates short SKUs such as P-4X9M and supports SKU search.",
                "Supports paginated inventory, search by name/barcode/SKU, categories, product images through R2, and base64 image handling.",
                "Supports manual add/update/delete, exact stock setting, cost-price updates, and product publication for storefronts.",
                "Fuzzy matching helps resolve misspelled products and protects against merging similar names incorrectly.",
                "Bulk AI upload parses up to 20 items at once for Tycoon users.",
                "Magic draft links help resolve ambiguous bulk restock items from the web.",
                "Low-stock thresholds and supplier phone numbers trigger restock alerts and supplier order links.",
            ],
            "Pitch: Merchants know what is in stock, what is low, what each item costs, and what can be shown online.",
        ),
        (
            "Customers, Loyalty, And Debtors",
            [
                "Customer records are unique per shop and include name, phone, royalty points, total spent, and last purchase.",
                "Owners can enable loyalty settings: points per purchase and redemption value per point.",
                "Sales can earn customer points, and customers can pay with points when enabled.",
                "Credit sales create or resolve debtor profiles and maintain balances.",
                "Debt payments can be recorded by WhatsApp or dashboard and are applied to outstanding debts.",
                "Debtor records include aliases, last product string, phone, due date, reminder status, and total debt.",
                "Debt due-date scheduler reminds debtors directly when a phone is saved, or reminds the owner if not.",
            ],
            "Pitch: TallyPadi helps owners remember who bought, who paid, who is owing, and who deserves rewards.",
        ),
        (
            "Expenses, Profit, And Reports",
            [
                "Records expenses with amount, description, category, date, and staff/owner actor context.",
                "Automatically adds new expense categories to the shop settings.",
                "Dashboard calculates total revenue, COGS, gross profit, total expenses, net profit, items sold, payment method totals, top products, debtors, and orders.",
                "WhatsApp reports cover sales, recent sales, stock, full summary, debts, expenses, best sellers, and sales comparison.",
                "Tycoon PDF reports can be generated for sales or full summaries when enabled.",
                "Generated PDF reports are cleaned up on a scheduled basis.",
            ],
            "Pitch: The owner stops guessing and starts seeing money in, money out, profit, debt, and stock movement.",
        ),
        (
            "Invoices And Bank Details",
            [
                "Tycoon users can create invoices from the dashboard or WhatsApp.",
                "Invoices include customer, items, quantities, unit price, totals, bank details snapshot, description, invoice number, and status.",
                "PDF invoices can be generated in A4 or thermal-friendly format.",
                "Marking an invoice as paid records a sale, updates daily stats, and deducts stock.",
                "Cancelling or deleting a paid invoice can undo the linked sale and restore stock.",
                "Bank details can be saved from settings or WhatsApp and are captured into invoices.",
            ],
            "Pitch: Merchants can send formal payment requests without designing invoices manually.",
        ),
        (
            "Orders And Job Tracking",
            [
                "Orders include description, customer name, optional phone, price, amount paid, balance, delivery date, and status.",
                "Statuses include pending, in progress, completed, delivered, and cancelled.",
                "Dashboard and WhatsApp support create, list, update, and cancel orders.",
                "Order reminders run before delivery dates and notify the owner.",
                "Balances make it useful for custom work, deposits, and pay-on-completion jobs.",
            ],
            "Pitch: TallyPadi is not only for walk-in sales. It can track jobs, deposits, and customer commitments.",
        ),
        (
            "Dashboard, Activity Feed, And Web POS",
            [
                "Dashboard returns shop profile, plan, wallet, subscription, currency, settings, storefront data, verification status, bank details, stats, inventory, sales, expenses, chart data, and top items.",
                "Main web navigation includes Dashboard, Activity, Sales, Orders, Expenses, Invoices, Product/Stocks, Customers, Debtors, Settings, Guide, Staff, Warehouse, Online Store, Ads Manager, and Subscription.",
                "Activity feed records important events such as expenses, low stock, subscriptions, wallet funding, and ads activity.",
                "Activity items can be filtered, counted, marked read, and resolved to the shop owner for staff accounts.",
                "PWA push notifications are available for owners, staff, support agents, and admin broadcasts.",
            ],
            "Pitch: The dashboard gives the owner the control room that WhatsApp alone cannot provide.",
        ),
        (
            "Online Store And Public Shop Link",
            [
                "Tycoon merchants can create a public shop slug at tallypadi.com/shop/{slug}.",
                "Storefront settings include business name, description, hero image, theme color, and location.",
                "Public shop pages show active Tycoon shops with published, in-stock products.",
                "Product pages include image, price, category, description, colors, sizes, stock status, SEO metadata, and seller verification badge when available.",
                "Public shop product list supports search, category filter, price sort, and pagination.",
                "Visit tracking increments daily shop visits so owners can see traffic.",
            ],
            "Pitch: The merchant gets a shareable catalog link without building a website.",
        ),
        (
            "Marketplace And Seller Verification",
            [
                "Marketplace lists products from active Tycoon sellers whose store setup is complete.",
                "Marketplace supports search, smart categories, category facets, state/city filters, sorting, boosted priority, product detail pages, and seller contact details.",
                "Marketplace SEO is generated from product, seller, price, category, location, and boost context.",
                "Seller verification supports NIN, national ID, driver's license, international passport, or government ID.",
                "Nigerian NIN flow can use government ID number and selfie; other IDs require document uploads and face captures.",
                "Admin reviews can approve, reject, delete, or request reverification.",
                "Verified sellers receive a visible Verified ID badge.",
            ],
            "Pitch: Marketplace plus verification helps honest sellers look credible and get discovered.",
        ),
        (
            "Ads Manager And Product Boosts",
            [
                "Only shop owners can fund the ads wallet and request boosts.",
                "Tycoon users can submit product boosts or custom campaigns.",
                "Promotion channels include Meta Ads, TikTok Ads, Google Ads, and TallyPadi Marketplace Boost.",
                f"Default boost plans are 3 days at {NAIRA}50,000, 7 days at {NAIRA}100,000, 14 days at {NAIRA}180,000, and 30 days at {NAIRA}300,000.",
                "Budget handling includes service fee, net campaign budget, safety reserve, optional FX buffer, and provider allocation.",
                "Campaigns go through TallyPadi admin review before fulfillment.",
                "System stores consent, policy checks, creative assets, Gemini ad suggestions, outbox events, notifications, runs, provider campaigns, and audit logs.",
                "Merchants can view campaigns, detail, metrics, top up, pause, stop, resume, and request changes.",
                "Admins can approve, reject, pause, resume, complete, update provider status, update metrics, refund, reallocate, or resubmit provider allocations.",
            ],
            "Pitch: The merchant funds a wallet and requests managed promotion. TallyPadi handles review, allocation, tracking, and provider workflow.",
        ),
        (
            "Staff, Permissions, And Ownership Controls",
            [
                "Tycoon active owners can add staff accounts, with a backend limit of 10 staff members.",
                "Staff receive WhatsApp invitations and log in through OTP sent to WhatsApp.",
                "Staff records connect to the owner and inherit shop settings.",
                "Staff can be updated or removed by the owner.",
                "Owner settings include permissions for dashboard, inventory, sales history, reports, customers, and settings.",
                "Staff activity can be restricted by permissions in the sidebar and APIs.",
                "Staff sales are recorded under the staff actor while inventory belongs to the owner shop.",
            ],
            "Pitch: Owners can delegate selling without giving staff full control of the business.",
        ),
        (
            "HQ, Warehouse, And Multi-Branch Operations",
            [
                "HQ users and HQ managers can list branches connected under one HQ account.",
                "HQ dashboard aggregates revenue, sales count, today's revenue, active branches, and recent network sales.",
                "Branch comparison ranks branches over the last seven days.",
                "HQ can create shop branches or warehouses.",
                "Stock transfer moves quantity from one branch to another and logs a transfer transaction.",
                "Staff can be promoted to HQ manager for branch-level operations.",
            ],
            "Pitch: TallyPadi can grow from one shop to branch and warehouse control.",
        ),
        (
            "Support, Customer Care, And Agent Tools",
            [
                "Support starts from WhatsApp when a user requests help.",
                "Support tickets move through queued, assigned, active, and closed states.",
                "Agents have login, status, ticket list, message threads, reply, pickup, close, delete, and escalation tools.",
                "Support can run in the dashboard and through WhatsApp agent routing.",
                "Tickets and messages are pushed over Socket.IO and Redis so agent screens update in real time.",
                "Support agents can view user deep dives, inventory, recent sales, staff, and message history, with destructive actions blocked at the support layer.",
            ],
            "Pitch: TallyPadi has built-in customer care infrastructure, not just merchant-facing features.",
        ),
        (
            "Admin, Investor, Blog, And Platform Operations",
            [
                "Admin analytics include user counts, active plans, active users, GMV, transaction count, and graph data.",
                "Admin user management supports deep dive, plan changes, subscription state changes, phone/email updates, messages, inventory actions, staff cleanup, and ads wallet top-up.",
                "Admin broadcast supports WhatsApp, push, and email, with targeting by plan, status, active users, trials, past due, or a specific identifier.",
                "Email templates support reusable HTML broadcasts with user placeholders.",
                "Admin ads review and marketplace verification review are first-class tabs in the web admin.",
                "Blog CMS supports drafts, publishing, SEO metadata, cover image, tags, content blocks, and public blog pages.",
                "Investor accounts can view high-level user and plan stats.",
            ],
            "Pitch: TallyPadi is built with the internal tools needed to run a serious SaaS operation.",
        ),
    ]

    add_heading(doc, "Full Feature Inventory", level=1)
    for heading, bullets, pitch in sections:
        add_heading(doc, heading, level=2)
        add_bullets(doc, bullets)
        add_callout(doc, "Marketer angle", pitch, fill="F7FAFC", accent=COLORS["dark_blue"])


def add_marketer_enablement(doc):
    doc.add_page_break()
    add_heading(doc, "How Marketers Should Explain It", level=1)
    add_table(
        doc,
        ["Customer Pain", "TallyPadi Message", "Feature Proof"],
        [
            ["I forget what I sold.", "Record sales instantly from WhatsApp or POS.", "Sales ledger, recent sales, reports, receipts, staff actor tracking."],
            ["Stock disappears.", "Stock reduces when sales are recorded and low-stock alerts can trigger.", "Inventory quantities, stock reports, low-stock thresholds, supplier messages."],
            ["Customers ask for receipts.", "Send professional PDF receipts and invoices.", "Receipt PDF, invoice PDF, bank detail snapshots."],
            ["People owe me.", "Track debtors, balances, due dates, and payments.", "Debtor model, debt payments, debt reminders."],
            ["I want to sell online.", "Publish a shop link and list products in a marketplace.", "Storefront, product pages, marketplace search, verified badge."],
            ["I need more buyers.", "Fund an ads wallet and request managed product boosts.", "Ads wallet, boost plans, review workflow, provider metrics."],
            ["My staff can make mistakes.", "Give staff controlled access and see who sold what.", "Staff OTP login, permissions, staff sale alerts."],
            ["I have branches.", "Track branches, compare sales, and transfer stock.", "HQ dashboard, branch compare, warehouse transfer."],
        ],
        [1.6, 2.35, 2.55],
    )

    add_heading(doc, "Demo Script", level=2)
    add_numbered(
        doc,
        [
            "Open WhatsApp and show a stock command: Add 20 cartons of Malt, cost 2000, selling 2500.",
            "Show the same product in the dashboard inventory with stock and price.",
            "Record a sale: Sold 3 Malt cash. Show stock reduced.",
            "Tap receipt and explain professional customer proof.",
            "Ask for report today. Show sales breakdown and profit when cost data exists.",
            "Create a debtor example: Sold 2 rice to Emeka on credit, due Friday. Then show debtor balance.",
            "Open the online store page and copy the public shop link.",
            "Show marketplace listing and verified seller badge if available.",
            "Open Ads Manager, explain wallet funding, boost plan, provider selection, and TallyPadi review.",
        ],
    )

    add_heading(doc, "Sample Phrases Marketers Can Use", level=2)
    add_bullets(
        doc,
        [
            "Your shop record is now inside WhatsApp, but with a full dashboard behind it.",
            "You can sell, restock, check debtors, and get reports without opening a spreadsheet.",
            "TallyPadi helps you look professional with receipts, invoices, and a public shop link.",
            "When you are ready to grow, you can publish products, get verified, and request ads from the same system.",
            "You can let staff sell without letting them control everything.",
        ],
    )

    add_heading(doc, "Objection Handling", level=2)
    add_table(
        doc,
        ["Objection", "Recommended Response"],
        [
            ["I already use a notebook.", "That works until you need totals, receipts, debt tracking, staff control, or online visibility. TallyPadi keeps the speed but removes manual counting."],
            ["I am not tech-savvy.", "Start with WhatsApp. The dashboard is there for reports, staff, stock, and online store setup."],
            ["Can my staff steal information?", "The product supports staff permissions and owner-controlled access. Staff can be limited based on what the owner allows."],
            ["Can customers buy directly online?", "The current storefront works as a catalog and buyer contact channel. Avoid promising full checkout or delivery automation unless it is added later."],
            ["Will ads guarantee sales?", "No platform should promise guaranteed results. TallyPadi helps package the product, budget, provider workflow, and reporting."],
            ["Will it replace my accountant?", "It gives sales, stock, expenses, receipts, and reports, but it is not a tax filing or regulated accounting service."],
        ],
        [1.75, 4.75],
    )


def add_plans_and_limits(doc):
    doc.add_page_break()
    add_heading(doc, "Plans, Packaging, And Claims", level=1)
    add_para(
        doc,
        "The codebase has several plan gates. The payment page and backend currently give the cleanest picture for marketers, while some older landing-page copy appears less precise. Use the table below as an internal guide, then confirm final packaging before creating public ads or sales sheets.",
    )
    add_table(
        doc,
        ["Plan Or Gate", "What The Code Shows", "Marketing Guidance"],
        [
            ["Free Trial", "User model defaults to 7-day trial. Landing copy lists WhatsApp assistant, PDF receipts, smart inventory, 1 user, and up to 50 products.", "Sell as a try-before-paying path, not as a permanent free tier unless leadership confirms it."],
            ["Oga Boss", f"Payment page price: {NAIRA}3,000 monthly, {NAIRA}15,000 for 6 months, {NAIRA}28,800 yearly. Payment page features core sales, basic inventory, daily profit summary, 1 user, standard support.", "Pitch as the core owner plan for basic business records."],
            ["Tycoon", f"Payment page price: {NAIRA}5,000 monthly, {NAIRA}27,000 for 6 months, {NAIRA}42,000 yearly. Backend gates staff, online store, invoices, bulk AI inventory upload, marketplace/ads boosts, and advanced reports to Tycoon in many places.", "Pitch as the growth plan for teams, online selling, invoices, advanced reporting, and promotion."],
            ["Staff limit", "Staff controller enforces a backend limit of 10 staff accounts for Tycoon owners. Some copy says unlimited or 5 in older places.", "Use 'up to 10 staff' until product leadership confirms another limit."],
            ["Ads budget", f"Default ads plans start at {NAIRA}50,000 and are separate from subscription fees.", "Tell customers ads require wallet funding and review; they are not included in monthly subscription."],
        ],
        [1.35, 3.25, 1.9],
    )

    add_heading(doc, "Do Not Overpromise", level=2)
    add_bullets(
        doc,
        [
            "Do not promise full ecommerce checkout, delivery routing, escrow, or customer payment collection on the public store. The current store is a product catalog/contact flow.",
            "Do not promise guaranteed ad results. The product supports managed campaign requests, review, provider workflow, and metrics.",
            "Do not promise tax filing, audited accounting, or legal compliance features.",
            "Do not promise every staff feature under Oga Boss. The backend heavily gates staff management to Tycoon.",
            "Do not promise fully automatic provider ads in every environment. Automation depends on TallyPadi-owned provider credentials and configuration.",
            "Do not promise bank reconciliation or inventory procurement automation beyond supplier restock message links.",
        ],
    )


def add_role_matrix(doc):
    add_heading(doc, "User Roles And What They See", level=1)
    add_table(
        doc,
        ["Role", "Primary Capabilities"],
        [
            ["Owner / Merchant", "Register, pay, manage shop, inventory, sales, reports, staff, store, verification, ads wallet, boosts, customers, debtors, invoices, orders, expenses."],
            ["Staff", "OTP login, record sales, handle orders and allowed dashboard areas based on owner permissions. Staff actions are recorded under staff identity."],
            ["HQ / HQ Manager", "View branches, dashboard, compare branches, transfer stock, create branches/warehouses, manage HQ-level operations."],
            ["Customer / Public Buyer", "Browse public shop and marketplace products, view seller location/contact, see verified seller badge when available."],
            ["Support Agent", "Log in, set status, pickup tickets, chat with users, escalate, close tickets, view support user details."],
            ["Admin", "Platform analytics, user management, broadcasts, global settings, ads review, seller verification review, support admin, blog CMS, investor accounts, wallet top-ups."],
            ["Investor", "View high-level registered, active plan, and trial user stats."],
        ],
        [1.55, 4.95],
    )


def add_onboarding_checklist(doc):
    add_heading(doc, "Onboarding Checklist For Marketers", level=1)
    add_bullets(
        doc,
        [
            "Phone number and email that the business owner controls.",
            "Business name, country, currency, and operating language.",
            "First product list with stock quantity, selling price, and cost price.",
            "Bank details if the merchant wants invoices.",
            "Customer/debtor list if the merchant sells on credit.",
            "Staff names and phone numbers if the owner wants team access.",
            "Storefront slug, description, location, theme color, hero image, and product images.",
            "Seller verification documents and selfie captures if the merchant wants a verified badge.",
            "Ads objective, budget, target location, target audience, creative notes, and product to boost.",
        ],
    )
    add_callout(
        doc,
        "Simple setup promise",
        "A strong onboarding promise is: start with one product, one sale, and one receipt. Once the owner sees that working, introduce inventory reports, debtors, storefront, and ads.",
        fill="F0F7F5",
        accent=COLORS["green"],
    )


def add_technical_appendix(doc):
    add_heading(doc, "Internal Technical Notes", level=1)
    add_para(
        doc,
        "This section is for marketers who need confidence about what powers the product, not for customer-facing pitches unless the buyer asks.",
    )
    add_table(
        doc,
        ["System Area", "Implementation Signal From Codebase"],
        [
            ["Backend", "Express + TypeScript API with MongoDB/Mongoose models and JWT authentication."],
            ["Web app", "Next.js dashboard, marketplace, shop pages, admin, support, investor, and SEO landing pages."],
            ["WhatsApp", "Meta WhatsApp Cloud API helpers for text, buttons, lists, flows, media, documents, CTA URLs, typing indicators, and read receipts."],
            ["AI", "Gemini parsing, prompt rules, ad SEO generation, fallback parsing, and structured intents."],
            ["Payments", "Paystack subscription payments and wallet funding with webhook signature checks and idempotent billing events."],
            ["Storage", "Cloudflare R2 presigned uploads and public product/verification/campaign asset handling."],
            ["Queues", "BullMQ and Redis queues for inbound messages, outbound replies, bulk messages, push notifications, and ads automation."],
            ["Realtime", "Socket.IO plus Redis pub/sub for support ticket and agent updates."],
            ["Schedulers", "Daily summaries, subscription expiry, debt reminders, order reminders, low stock alerts, PDF cleanup, and ads maintenance."],
            ["Security", "Helmet, CORS policy, rate limits, WhatsApp and Paystack signature handling, input validation, zod schemas, sanitization, and anti-duplicate message processing."],
        ],
        [1.55, 4.95],
    )

    add_heading(doc, "Source Confidence", level=2)
    add_para(
        doc,
        "This guide was drafted from the repository files in the current workspace, including API route registration, controllers, services, models, web pages, and existing API documentation. It reflects code-level capabilities visible on May 22, 2026.",
        italic=True,
        color=COLORS["muted"],
    )


def build():
    doc = Document()
    setup_document(doc)
    add_cover(doc)
    add_executive_summary(doc)
    add_customer_journey(doc)
    add_feature_sections(doc)
    add_marketer_enablement(doc)
    add_plans_and_limits(doc)
    add_role_matrix(doc)
    add_onboarding_checklist(doc)
    add_technical_appendix(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
